from docx import Document
from docx.shared import Inches

# สร้างเอกสาร Word ใหม่
doc = Document()

# เพิ่มข้อความลงในเอกสาร
doc.add_heading('สวัสดีชาวโลก', level=1)
doc.add_paragraph('นี่คือเอกสารที่สร้างขึ้นโดยใช้ Python และไลบรารี docx')

# เพิ่มรูปภาพลงในเอกสาร
doc.add_heading('รูปภาพที่น่ารัก', level=2)
doc.add_picture('cute_cat.jpg', width=Inches(4.0))

# เพิ่มตารางลงในเอกสาร
doc.add_heading('ตารางข้อมูล', level=2)
table = doc.add_table(rows=3, cols=3)
table.cell(0, 0).text = 'ชื่อ'
table.cell(0, 1).text = 'อายุ'
table.cell(0, 2).text = 'อาชีพ'
table.cell(1, 0).text = 'สมชาย'
table.cell(1, 1).text = '30'
table.cell(1, 2).text = 'วิศวกร'
table.cell(2, 0).text = 'สมหญิง'
table.cell(2, 1).text = '28'
table.cell(2, 2).text = 'นักเรียน'

# บันทึกเอกสารลงในไฟล์
doc.save('example.docx')